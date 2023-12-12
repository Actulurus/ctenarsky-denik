import requests, docx
import threading, dotenv, time, sys, os

from docx.shared import Pt

import text_utils, gpt

file_path = os.path.abspath(os.path.dirname(__file__))

if not os.path.exists(os.path.join(file_path, '.env')):
    with open(os.path.join(file_path, ".env"), 'w') as f:
        f.write("""SERVER_KEY=YOUR_KEY_HERE
SERVER_URL='http://getkey.stuckinvim.com/api/data?api_key='
OPENAI_KEY=YOUR_KEY_HERE_only_use_this_if_you_dont_have_a_server_key
TESTMODE=false""")

        print("Missing .env file, created one for you.")
        print("Please enter your StuckInVim key in the .env file")
        exit()

dotenv.load_dotenv()

SERVER_URL = os.getenv("SERVER_URL")
SERVER_KEY = os.getenv("SERVER_KEY")
OPENAI_KEY = os.getenv("OPENAI_KEY")

if not SERVER_KEY and not OPENAI_KEY:
    print("Missing API key. Please set the SERVER_KEY or OPENAI_KEY environment variable.")
    exit(1)

API_KEY = OPENAI_KEY

if not API_KEY and SERVER_URL:
    result = requests.get(SERVER_URL + SERVER_KEY).json()

    if result.get("status") == "200":
        API_KEY = result["key"]
    else:
        raise Exception("Failed to get API key from server api API.")
elif not API_KEY and not SERVER_URL:
    raise Exception("Missing server URL.")

if not os.path.exists(os.path.join(file_path, "output")):
    os.mkdir(os.path.join(file_path, "output"))

gpt = gpt.init(API_KEY)

def set_document_format(document, line_spacing=None, font_size=None, font_name=None, alignment=None):
    # Set line spacing for the entire document
    if line_spacing is not None:
        document.styles['Normal'].paragraph_format.line_spacing = line_spacing

    # Set font size and font name for the entire document
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if font_size is not None:
                run.font.size = Pt(font_size)
            if font_name is not None:
                run.font.name = font_name

    # Set alignment for the entire document
    if alignment is not None:
        for paragraph in document.paragraphs:
            paragraph.alignment = alignment

def write_docx(text):
    doc = docx.Document()

    doc.add_paragraph(text)
    set_document_format(doc, line_spacing=1.15, font_size=12, font_name='Times New Roman')

    newline = "\n"
    doc.save(os.path.join(file_path, "output", f"{text.split(newline)[0]}.docx"))

def run(filters=[]):
    links = text_utils.extract_links(filters=filters)

    for link in links:
        text = text_utils.extract_text_from_url(link)
        if text:
            def run_gpt():
                result = gpt["complete"](text)

                if result:
                    print(result)
                    write_docx(result)
                else:
                    print("Failed to run GPT on text from " + link)

            threading.Thread(target=run_gpt).start()

            time.sleep(10)
        else:
            print("Failed to extract text from " + link)

        if os.getenv("TESTMODE") == "true":
            break

if __name__ == "__main__":
    filters = []
    if len(sys.argv) > 1:
        filters = sys.argv[1:]
    
    run(filters=filters)